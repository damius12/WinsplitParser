import os

import docx


def _set_margins(doc: docx.Document) -> None:
    section = doc.sections[0]
    section.top_margin = docx.shared.Cm(1.5)
    section.bottom_margin = docx.shared.Cm(1.5)
    section.left_margin = docx.shared.Cm(1.5)
    section.right_margin = docx.shared.Cm(1.5)


def _set_font(doc: docx.Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = docx.shared.Pt(11)


def create_word(
    event_title: str, results_text: str, image_path: str, output_path: str
) -> None:
    doc = docx.Document()
    _set_margins(doc)
    _set_font(doc)
    doc.add_heading(event_title, 1)
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=docx.shared.Cm(12))
    doc.add_paragraph(results_text)
    doc.save(output_path)
